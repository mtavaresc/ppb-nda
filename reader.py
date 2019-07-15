import sys
from docx import Document
from json import dumps
from re import findall, sub, match
# from itertools import groupby
from base import Session, engine
from models import *


def create_or_drop_db(drop=False):
    if drop:
        Base.metadata.drop_all(engine)
    Base.metadata.create_all(engine)


def extract_data_from_playbook(doc):
    issues, principles_paragraph = [], []
    num_paragraph = 0
    for para in doc.paragraphs:
        num_paragraph += 1
        if para.text.lower().startswith('Issue'.lower()):
            issue_paragraph = para.text.split('.')[0]
            if match(r'(Issue\s[#]\d+[A-Z])', para.text) is not None:
                issue_id = findall(r'(Issue\s[#]\d+[A-Z])', issue_paragraph)[0]
                issue_id = issue_id[issue_id.index('#') + 1:]
                if sub(r'(Issue\s[#]\d+[A-Z])', '', issue_paragraph).split(':')[1].strip() == 'Customer’s request':
                    issue_name = sub(r'(Issue\s[#]\d+[A-Z])', '', issue_paragraph).split(':')[2].strip()
                else:
                    issue_name = sub(r'(Issue\s[#]\d+[A-Z])', '', issue_paragraph).split(':')[1].strip()
                if issues[-1][0] == issue_id[:-1]:
                    _dict = {
                        'id': issue_id[-1:],
                        'customer_request': issue_name
                    }

                    tup = issues[-1] + (_dict,)
                    issues[issues.index(issues[-1])] = tup
            else:
                issue_id = issue_paragraph[issue_paragraph.index('#') + 1:issue_paragraph.index(':')]
                issue_name = issue_paragraph.split(':')[1].strip()
                issues.append((issue_id, issue_name))
        elif para.text.lower().startswith('Customer’s request'.lower()):
            cr = para.text.split(':')[1].strip()
            _dict = {
                'id': 'A',
                'customer_request': cr
            }
            tup = issues[-1] + (_dict,)
            issues.pop()
            issues.append(tup)

        if len(issues) == 0:
            continue

        if len(issues[-1]) == 3 or len(issues[-1]) == 4:
            found = False
            num_paragraph_2 = 0
            for p in doc.paragraphs:
                num_paragraph_2 += 1
                if num_paragraph > num_paragraph_2:
                    continue
                next_topic = ['Sample Language', 'Dell EMC Standard Language', 'Fallback', 'Approval']
                if p.text.lower().startswith('Principle'.lower()):
                    if len(p.text.split(':')[1].strip()) > 0:
                        found = True

                if found:
                    if any(p.text.lower().startswith(word.lower()) for word in next_topic) is False and p.text != '':
                        principles_paragraph.append(sub(r'(Principle+?)(s\b|\b)[:]', '', p.text).strip())

                if any(p.text.lower().startswith(word.lower()) for word in next_topic):
                    break

            if len(principles_paragraph) > 0:
                principles = ''.join(principles_paragraph)
                tup = issues[-1] + (principles,)
                issues.pop()
                issues.append(tup)
                principles_paragraph.clear()

    return issues


def insert_db(issues):
    s = Session()
    for i in issues:
        issue = i[0]
        name = i[1]
        if isinstance(i[-1], dict):
            principle = i[-2]
        else:
            principle = i[-1]
        print(principle)
        # s.merge(Issues(issue, name))
        # for idx in range(2, len(i)):
        #     if isinstance(i[idx], dict):
        #         s.merge(CustomersRequest(issue, i[idx]['id'], i[idx]['customer_request'], principle))
        #         # try:
        #         # except Exception as e:
        #         #     print(f'#1 Error: {e}')
        #         #     return False
        #
        # s.commit()
    return True


if __name__ == '__main__':
    # stdoutOrigin = sys.stdout
    # sys.stdout = open('log.txt', 'w')

    document = Document('samples/playbook_nda.docx')
    # create_or_drop_db()
    rs = extract_data_from_playbook(document)
    # print(dumps(rs, indent=4))
    insert_db(rs)

    # sys.stdout.close()
    # sys.stdout = stdoutOrigin
